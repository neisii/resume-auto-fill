import copy
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

from .formatter import format_value
from .loader import flatten_profile
from .matcher import FieldMatcher

_INDEXED_KEY_RE = re.compile(r"^([a-z]+)\.(\d+)\.")

# Horizontal table is only filled when at least one of these fields appears in col_map
_IDENTIFYING_FIELDS = {
    "projects.0.name",
    "projects.0.role",
    "projects.0.description",
    "projects.0.client",
    "projects.0.company",
    "projects.0.environment.os",
    "career.0.company",
    "career.0.duties",
    "career.0.period",
    "education.0.school",
    "education.0.major",
    "education.0.end",
    "certifications.0.name",
    "certifications.0.date",
}


def fill_template(
    template_path: str | Path,
    profile: dict,
    matcher: FieldMatcher,
    output_path: str | Path,
    overwrite: set[str] | None = None,
) -> dict[str, int]:
    doc = Document(str(template_path))
    flat = flatten_profile(profile)
    stats = {"filled": 0, "skipped": 0}
    ow = overwrite or set()

    for table in doc.tables:
        rows_cells = [[cell for cell in row.cells] for row in table.rows]
        if _count_table_labels(rows_cells, matcher) < 2:
            continue
        col_map = _build_col_map(rows_cells, matcher)
        if _is_horizontal_table(rows_cells, col_map):
            header_count = _count_header_rows(rows_cells)
            _fill_horizontal(rows_cells[header_count:], col_map, flat, stats, ow)
        else:
            _fill_vertical(rows_cells, flat, matcher, stats, ow)

    doc.save(str(output_path))
    return stats


# ── helpers ──────────────────────────────────────────────────────────────────

def _first_line(text: str) -> str:
    """Use only the first line of a cell for label matching."""
    return text.split("\n")[0].strip()


def _count_table_labels(rows_cells: list[list[_Cell]], matcher: FieldMatcher) -> int:
    """Count distinct label cells across the entire table."""
    seen: set[int] = set()
    count = 0
    for row_cells in rows_cells:
        for _, cell in _unique_cells(row_cells, seen):
            if matcher.match(_first_line(cell.text)) is not None:
                count += 1
    return count


def _unique_cells(cells: list[_Cell], seen: set) -> list[tuple[int, _Cell]]:
    """Yield (col_idx, cell) for cells not already in seen (by _tc id)."""
    result = []
    for c_idx, cell in enumerate(cells):
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            result.append((c_idx, cell))
    return result


# ── placeholder detection ─────────────────────────────────────────────────────

_PLACEHOLDER_PATTERNS = [
    re.compile(r'^y{2,4}[./년]m{1,2}([./월]d{1,2}일?)?$', re.IGNORECASE),  # yyyy.mm, yyyy.mm.dd
    re.compile(r'^\d*0{3,}[-./]\d*0{2,}'),   # 0000-00, 0000.00 형식
    re.compile(r'^x+[-x\s]+x+$', re.IGNORECASE),   # xxx-xxxx-xxxx
    re.compile(r'입력(하세요|바랍니다|해주세요)?'),
    re.compile(r'기재(하세요|바랍니다|해주세요)?'),
    re.compile(r'작성(하세요|바랍니다|해주세요)?'),
    re.compile(r'^예\s*\)'),
]


def is_placeholder(text: str, overwrite: set[str]) -> bool:
    """Return True if the cell text is a known placeholder that should be overwritten."""
    t = text.strip()
    if not t:
        return True
    if t in overwrite:
        return True
    for pattern in _PLACEHOLDER_PATTERNS:
        if pattern.search(t):
            return True
    return False


def _row_is_fillable(cells: list[_Cell], overwrite: set[str], seen: set | None = None) -> bool:
    """Return True if every unique cell is empty or a placeholder."""
    if seen is None:
        seen = set()
    for _, cell in _unique_cells(cells, seen):
        if not is_placeholder(cell.text, overwrite):
            return False
    return True


def _row_is_empty(cells: list[_Cell], seen: set | None = None) -> bool:
    if seen is None:
        seen = set()
    for _, cell in _unique_cells(cells, seen):
        if cell.text.strip():
            return False
    return True


# ── horizontal table handling ─────────────────────────────────────────────────

def _build_col_map(rows_cells: list[list[_Cell]], matcher: FieldMatcher) -> dict[int, str]:
    """Scan first 2 rows and map col_idx → field_key for any matching labels."""
    seen: set[int] = set()
    col_map: dict[int, str] = {}
    for row_cells in rows_cells[:2]:
        for c_idx, cell in _unique_cells(row_cells, seen):
            fk = matcher.match(_first_line(cell.text))
            if fk is not None and c_idx not in col_map:
                col_map[c_idx] = fk
    return col_map


def _is_horizontal_table(rows_cells: list[list[_Cell]], col_map: dict[int, str]) -> bool:
    """
    Horizontal table: ≥ 3 distinct label columns in first 2 rows AND
    at least one identifying field AND at least one all-empty data row.
    """
    if len(col_map) < 3:
        return False
    if not any(v in _IDENTIFYING_FIELDS for v in col_map.values()):
        return False
    for row_cells in rows_cells[2:]:
        if _row_is_empty(row_cells):
            return True
    return False


def _count_header_rows(rows_cells: list[list[_Cell]]) -> int:
    """Count leading rows where every unique cell has text (no empty cells)."""
    for row_idx, cells in enumerate(rows_cells):
        seen: set[int] = set()
        for _, cell in _unique_cells(cells, seen):
            if not cell.text.strip():
                return row_idx
    return len(rows_cells)


def _fill_horizontal(
    data_rows: list[list[_Cell]],
    col_map: dict[int, str],
    flat: dict[str, Any],
    stats: dict,
    overwrite: set[str] | None = None,
) -> None:
    ow = overwrite or set()
    section = _detect_section(col_map)
    if not section:
        return
    max_entries = _section_count(flat, section)
    entry_idx = 0

    for row_cells in data_rows:
        if entry_idx >= max_entries:
            break
        if not _row_is_fillable(row_cells, ow):
            continue

        for c_idx, base_key in col_map.items():
            if c_idx >= len(row_cells):
                continue
            cell = row_cells[c_idx]
            eff_key = _reindex_key(base_key, entry_idx)
            # Skip fields that belong to a different section than the detected table section
            m = _INDEXED_KEY_RE.match(eff_key)
            if m and m.group(1) != section:
                stats["skipped"] += 1
                continue
            value = flat.get(eff_key)
            if value is None:
                stats["skipped"] += 1
                continue
            formatted = format_value(value, eff_key)
            if formatted:
                _set_cell_text(cell, formatted)
                stats["filled"] += 1

        entry_idx += 1


def _detect_section(col_map: dict[int, str]) -> str | None:
    """Return the profile section name (e.g. 'career') from the first indexed field_key."""
    for fk in col_map.values():
        m = _INDEXED_KEY_RE.match(fk)
        if m:
            return m.group(1)
    return None


def _reindex_key(key: str, idx: int) -> str:
    """Replace the numeric index in any section key: career.0.x → career.N.x"""
    return _INDEXED_KEY_RE.sub(lambda m: f"{m.group(1)}.{idx}.", key)


def _section_count(flat: dict, section: str) -> int:
    """Count how many entries exist for a given section in the flat profile."""
    max_idx = -1
    for k in flat:
        m = _INDEXED_KEY_RE.match(k)
        if m and m.group(1) == section:
            max_idx = max(max_idx, int(m.group(2)))
    return max_idx + 1 if max_idx >= 0 else 0


# ── vertical table handling ───────────────────────────────────────────────────

def _label_count_in_row(cells: list[_Cell], matcher: FieldMatcher) -> int:
    """Count unique cells that match a label (for header-row detection)."""
    seen: set[int] = set()
    count = 0
    for _, cell in _unique_cells(cells, seen):
        if matcher.match(_first_line(cell.text)) is not None:
            count += 1
    return count


def _fill_vertical(
    rows_cells: list[list[_Cell]],
    flat: dict[str, Any],
    matcher: FieldMatcher,
    stats: dict,
    overwrite: set[str] | None = None,
) -> None:
    ow = overwrite or set()
    visited: set[int] = set()

    for row_idx, cells in enumerate(rows_cells):
        # Rows with ≥ 3 label cells are column-header rows, not key-value rows.
        # Treating their labels as row-labels would fill adjacent/below cells incorrectly.
        if _label_count_in_row(cells, matcher) >= 3:
            continue

        for col_idx, cell in enumerate(cells):
            cell_id = id(cell._tc)
            if cell_id in visited:
                continue
            visited.add(cell_id)

            label = _first_line(cell.text)
            field_key = matcher.match(label)
            if field_key is None:
                continue

            value_cell = _find_value_cell(rows_cells, row_idx, col_idx, visited, matcher, ow)
            if value_cell is None:
                stats["skipped"] += 1
                continue

            value = flat.get(field_key)
            if value is None:
                stats["skipped"] += 1
                continue

            formatted = format_value(value, field_key)
            if formatted:
                _set_cell_text(value_cell, formatted)
                visited.add(id(value_cell._tc))
                stats["filled"] += 1


def _find_value_cell(
    rows_cells: list[list[_Cell]],
    row_idx: int,
    col_idx: int,
    visited: set,
    matcher: FieldMatcher,
    overwrite: set[str] | None = None,
) -> _Cell | None:
    ow = overwrite or set()
    cells = rows_cells[row_idx]
    current_tc_id = id(cells[col_idx]._tc)

    for next_col in range(col_idx + 1, len(cells)):
        candidate = cells[next_col]
        cid = id(candidate._tc)
        if cid == current_tc_id:
            continue
        if cid in visited:
            continue
        candidate_text = candidate.text.strip()
        if candidate_text:
            if matcher.match(_first_line(candidate_text)) is not None:
                break  # It's a label — stop scanning
            if is_placeholder(candidate_text, ow):
                return candidate  # Placeholder — overwrite
            continue  # Meaningful content — skip
        return candidate  # Empty cell

    # Fallback: cell directly below (empty or placeholder)
    if row_idx + 1 < len(rows_cells) and col_idx < len(rows_cells[row_idx + 1]):
        below = rows_cells[row_idx + 1][col_idx]
        bid = id(below._tc)
        if bid not in visited and bid != current_tc_id:
            if is_placeholder(below.text, ow):
                return below

    return None


# ── cell text writing ─────────────────────────────────────────────────────────

def _set_cell_text(cell: _Cell, text: str) -> None:
    if not cell.paragraphs:
        return

    para = cell.paragraphs[0]

    existing_rpr = None
    if para.runs:
        rpr_elem = para.runs[0]._element.find(qn("w:rPr"))
        if rpr_elem is not None:
            existing_rpr = copy.deepcopy(rpr_elem)

    for r_elem in para._element.findall(qn("w:r")):
        para._element.remove(r_elem)

    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            br_run = OxmlElement("w:r")
            if existing_rpr is not None:
                br_run.append(copy.deepcopy(existing_rpr))
            br_run.append(OxmlElement("w:br"))
            para._element.append(br_run)

        if line:
            r_elem = OxmlElement("w:r")
            if existing_rpr is not None:
                r_elem.append(copy.deepcopy(existing_rpr))
            t_elem = OxmlElement("w:t")
            t_elem.text = line
            if line.startswith(" ") or line.endswith(" "):
                t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r_elem.append(t_elem)
            para._element.append(r_elem)
